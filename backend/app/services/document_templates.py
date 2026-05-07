from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "templates" / "analysis"
NOTICE_TEMPLATE = TEMPLATE_DIR / "tax_notice_template.docx"
REPORT_TEMPLATE = TEMPLATE_DIR / "officer_report_template.docx"


def _chinese_date(value: datetime | None = None) -> str:
    value = value or datetime.now()
    digits = "〇一二三四五六七八九"
    year = "".join(digits[int(ch)] for ch in str(value.year))
    month_map = {
        1: "一月", 2: "二月", 3: "三月", 4: "四月", 5: "五月", 6: "六月",
        7: "七月", 8: "八月", 9: "九月", 10: "十月", 11: "十一月", 12: "十二月",
    }
    if value.day <= 10:
        day = "十日" if value.day == 10 else f"{digits[value.day]}日"
    elif value.day < 20:
        day = f"十{digits[value.day % 10]}日"
    elif value.day in (20, 30):
        day = f"{digits[value.day // 10]}十日"
    else:
        day = f"{digits[value.day // 10]}十{digits[value.day % 10]}日"
    return f"{year}年{month_map[value.month]}{day}"


def _document_date_text(value: str | None = None) -> str:
    text = str(value or "").strip()
    if not text:
        return _chinese_date()
    for fmt in ("%Y-%m-%d", "%Y/%m/%d"):
        try:
            return _chinese_date(datetime.strptime(text, fmt))
        except ValueError:
            continue
    return text


def _clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _set_run_font(paragraph, size: int = 16, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "仿宋"
        run.font.size = Pt(size)
        run.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")


def qn(tag: str):
    from docx.oxml.ns import qn as _qn
    return _qn(tag)


def _add_paragraph(doc: Document, text: str = "", *, align=None, size: int = 16, bold: bool = False, first_indent: bool = False):
    paragraph = doc.add_paragraph(text)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.25
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(size * 2)
    _set_run_font(paragraph, size=size, bold=bold)
    return paragraph


def _safe_join(items: list[Any]) -> str:
    return "；".join(str(item) for item in items if str(item).strip()) or "暂无"


def _source_refs_text(refs: list[dict[str, Any]]) -> str:
    parts = []
    for ref in refs or []:
        label = ref.get("dataset_label") or ref.get("dataset_kind") or "资料"
        period = ref.get("period") or "期间待核实"
        field = ref.get("field_label") or ref.get("field_name") or "字段"
        value = ref.get("value", "")
        parts.append(f"{label}[{period}] {field}={value}")
    return "；".join(parts) or "暂无"


def render_notice_docx(payload: dict[str, Any], agency_name: str = "主管税务机关") -> bytes:
    doc = Document(str(NOTICE_TEMPLATE))
    _clear_document_body(doc)

    agency_name = payload.get("agency_name") or agency_name
    _add_paragraph(doc, agency_name, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True)
    _add_paragraph(doc, "税务事项通知书", align=WD_ALIGN_PARAGRAPH.CENTER, size=26, bold=True)
    notice_no = payload.get("document_number") or f"税通〔{datetime.now().year}〕{payload.get('task_id', '0000')[-6:]}号"
    _add_paragraph(doc, notice_no, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    _add_paragraph(doc)

    enterprise = payload.get("enterprise_name") or "待补充企业名称"
    taxpayer_id = payload.get("taxpayer_id") or "待补充纳税人识别号"
    _add_paragraph(doc, f"{enterprise}（纳税人识别号：{taxpayer_id}）：", size=16)

    issues = payload.get("issues") or []
    reason = "涉税风险核实整改" if issues else "涉税资料补充核实"
    basis = "根据《中华人民共和国税收征收管理法》第二十五条、第六十四条等有关规定，以及你单位申报、发票、财务报表等资料比对情况。"
    _add_paragraph(doc, f"事由：{reason}", size=16, first_indent=True)
    _add_paragraph(doc, f"依据：{basis}", size=16, first_indent=True)
    _add_paragraph(doc, "通知内容：经案头分析，你单位相关涉税数据存在需进一步核实的事项。现将有关事项通知如下：", size=16, first_indent=True)

    if not issues:
        _add_paragraph(doc, "一、当前未识别出明确风险事项，请补充完整发票、申报和财务报表资料后配合进一步核实。", size=16, first_indent=True)
    for index, issue in enumerate(issues, start=1):
        _add_paragraph(doc, f"{index}. {issue.get('risk_type', '涉税风险')}：{issue.get('issue', '需进一步核实')}", size=16, first_indent=True)
        _add_paragraph(doc, f"涉及期间：{issue.get('period', '待核实')}", size=16, first_indent=True)
        _add_paragraph(doc, f"规则名称：{issue.get('rule_name', '规则待补充')}", size=16, first_indent=True)
        _add_paragraph(doc, f"触发原因：{issue.get('trigger_reason', issue.get('issue', '需进一步核实'))}", size=16, first_indent=True)
        _add_paragraph(doc, f"涉及数据：{_source_refs_text(issue.get('source_data_refs', []))}", size=16, first_indent=True)
        _add_paragraph(doc, f"计算过程：{issue.get('calculation_text', '未生成计算说明')}", size=16, first_indent=True)
        _add_paragraph(doc, f"判断阈值：{issue.get('threshold_text', '按申报、发票、财务资料一致性综合判断')}", size=16, first_indent=True)
        _add_paragraph(doc, f"依据数据：{issue.get('basis_data', '申报、发票、财务数据比对结果')}", size=16, first_indent=True)
        _add_paragraph(doc, f"证据要点：{_safe_join(issue.get('evidence', []))}", size=16, first_indent=True)
        _add_paragraph(doc, f"整改要求：{issue.get('rectify_advice', '请说明原因并提供相关佐证资料。')}", size=16, first_indent=True)

    _add_paragraph(doc, f"请你单位于{payload.get('rectification_deadline', '收到通知后5个工作日内')}完成情况说明、资料补正或整改反馈。", size=16, first_indent=True)
    _add_paragraph(doc, f"联系人：{payload.get('contact_person', '主管税务人员')}；联系电话：{payload.get('contact_phone', '请在系统配置中补充联系人电话')}。", size=16, first_indent=True)
    _add_paragraph(doc)
    _add_paragraph(doc, agency_name, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    _add_paragraph(doc, _document_date_text(payload.get("document_date")), align=WD_ALIGN_PARAGRAPH.CENTER, size=16)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_officer_report_docx(payload: dict[str, Any]) -> bytes:
    doc = Document(str(REPORT_TEMPLATE))
    _clear_document_body(doc)

    agency_name = payload.get("agency_name", "")
    document_number = payload.get("document_number") or payload.get("task_id", "待生成")
    _add_paragraph(doc, "税务疑点核实报告", align=WD_ALIGN_PARAGRAPH.CENTER, size=24, bold=True)
    _add_paragraph(doc, f"报告编号：{document_number}", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    if agency_name:
        _add_paragraph(doc, f"出具机关：{agency_name}", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _add_paragraph(doc, f"文书日期：{payload.get('document_date') or datetime.now().strftime('%Y-%m-%d')}", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _add_paragraph(doc)

    _add_paragraph(doc, "一、企业基本信息", size=16, bold=True)
    _add_paragraph(doc, f"纳税人识别号：{payload.get('taxpayer_id', '待补充')}", size=16)
    _add_paragraph(doc, f"企业名称：{payload.get('enterprise_name', '待补充')}", size=16)
    _add_paragraph(doc, f"分析任务：{payload.get('task_name', '案头分析任务')}", size=16)
    _add_paragraph(doc, f"风险数量：{payload.get('risk_count', 0)}", size=16)

    data_summary = payload.get("data_summary") or {}
    _add_paragraph(doc, "二、数据概况", size=16, bold=True)
    _add_paragraph(doc, f"分析期间：{', '.join(data_summary.get('periods') or []) or '待核实'}", size=16)
    _add_paragraph(doc, f"销项发票数：{data_summary.get('sales_invoice_count', 0)}；进项发票数：{data_summary.get('purchase_invoice_count', 0)}；增值税申报期间数：{data_summary.get('vat_return_periods', 0)}；企业所得税期间数：{data_summary.get('cit_return_periods', 0)}；财务报表期间数：{data_summary.get('financial_statement_periods', 0)}。", size=16, first_indent=True)

    risks = payload.get("risks") or []
    _add_paragraph(doc, "三、疑点核实详情", size=16, bold=True)
    if not risks:
        _add_paragraph(doc, "当前未识别出明确风险，建议补充完整资料后再次开展案头分析。", size=16, first_indent=True)
    for index, risk in enumerate(risks, start=1):
        _add_paragraph(doc, f"疑点{index}：{risk.get('risk_type', '涉税风险')}", size=16, bold=True)
        _add_paragraph(doc, f"数据结论：{risk.get('issue', '需进一步核实')}", size=16, first_indent=True)
        _add_paragraph(doc, f"涉及期间：{risk.get('period', '待核实')}；风险等级：{risk.get('severity', '待核实')}；可信度：{risk.get('confidence', '待核实')}", size=16, first_indent=True)
        _add_paragraph(doc, f"规则名称：{risk.get('rule_name', '规则待补充')}", size=16, first_indent=True)
        _add_paragraph(doc, f"触发原因：{risk.get('trigger_reason', risk.get('issue', '需进一步核实'))}", size=16, first_indent=True)
        _add_paragraph(doc, f"涉及数据：{_source_refs_text(risk.get('source_data_refs', []))}", size=16, first_indent=True)
        _add_paragraph(doc, f"计算过程：{risk.get('calculation_text', '未生成计算说明')}", size=16, first_indent=True)
        _add_paragraph(doc, f"判断阈值：{risk.get('threshold_text', '按申报、发票、财务资料一致性综合判断')}", size=16, first_indent=True)
        _add_paragraph(doc, f"证据要点：{_safe_join(risk.get('evidence', []))}", size=16, first_indent=True)
        _add_paragraph(doc, f"核实方向：{risk.get('verification_focus', '请结合申报、发票、账簿和资金流水进一步核实。')}", size=16, first_indent=True)
        _add_paragraph(doc, f"应要求企业提供资料：{_safe_join(risk.get('required_materials', []))}", size=16, first_indent=True)
        _add_paragraph(doc, f"判断标准：{risk.get('judgment_rule', '结合业务真实性、资料完整性和申报一致性判断。')}", size=16, first_indent=True)

    warnings = data_summary.get("data_warnings") or []
    _add_paragraph(doc, "四、资料完整性提醒", size=16, bold=True)
    if warnings:
        for item in warnings:
            _add_paragraph(doc, f"• {item}", size=16, first_indent=True)
    else:
        _add_paragraph(doc, "当前未发现明确资料完整性提醒。", size=16, first_indent=True)

    _add_paragraph(doc, "五、核实工作建议", size=16, bold=True)
    suggestions = [
        "优先核实风险等级较高、涉及金额较大或证据链较完整的疑点。",
        "要求企业按疑点逐项提供合同、发票、账簿、记账凭证、资金流水、库存和物流等佐证资料。",
        "核实过程中应保留资料调取、企业说明、比对口径和判断依据。",
        "发现已排除、整改中或已整改事项，应同步记入风险记录台账。",
    ]
    for item in suggestions:
        _add_paragraph(doc, item, size=16, first_indent=True)

    _add_paragraph(doc, "六、附件清单", size=16, bold=True)
    _add_paragraph(doc, "1. 税务事项通知书", size=16)
    _add_paragraph(doc, "2. 案头分析资料清单", size=16)
    _add_paragraph(doc, "3. 发票、申报和财务数据比对明细", size=16)
    _add_paragraph(doc)
    _add_paragraph(doc, "报告生成人：税务案头分析系统", size=16)
    _add_paragraph(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", size=16)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
